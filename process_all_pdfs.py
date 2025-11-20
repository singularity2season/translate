import requests
import os
import glob
import xml.etree.ElementTree as ET
import time
from dotenv import load_dotenv
from docx import Document # 追加: Word作成用
from docx.shared import Pt # 追加: フォントサイズ指定用

# --- ▼ 設定エリア ▼ ---

load_dotenv()

DEEPL_API_KEY = os.getenv("DEEPL_API_KEY")
DEEPL_URL = os.getenv("DEEPL_API_URL", "https://api-free.deepl.com/v2/translate")
GROBID_URL = os.getenv("GROBID_API_URL", "http://localhost:8070/api/processFulltextDocument")

TARGET_LANG = "JA"
GROBID_TIMEOUT = 180
INPUT_DIR = "input_pdf"
OUTPUT_DIR = "output_pdf"

# 出力フォルダ設定
OUTPUT_XML_DIR = os.path.join(OUTPUT_DIR, "xml")
OUTPUT_DOCX_DIR = os.path.join(OUTPUT_DIR, "docx") # 変更: docx用フォルダ
NAMESPACES = {'tei': 'http://www.tei-c.org/ns/1.0'}

# ------------------------------------------------

def setup_directories():
    os.makedirs(OUTPUT_XML_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DOCX_DIR, exist_ok=True)

def translate_text_via_deepl(text):
    """DeepL APIを使ってテキストを翻訳"""
    if not text or not text.strip():
        return ""
    
    params = {
        "auth_key": DEEPL_API_KEY,
        "text": text,
        "target_lang": TARGET_LANG
    }

    try:
        response = requests.post(DEEPL_URL, data=params, timeout=30)
        if response.status_code == 200:
            return response.json()["translations"][0]["text"]
        elif response.status_code == 456:
            return "[Translation Error: Quota Exceeded]"
        else:
            print(f"  ⚠️ DeepLエラー: {response.status_code}")
            return text # エラー時は原文を返す
    except Exception as e:
        print(f"  ⚠️ 通信エラー: {e}")
        return text

def translate_long_text(full_text):
    """長文を段落ごとに翻訳して結合"""
    paragraphs = full_text.split("\n\n")
    translated_paragraphs = []
    
    print(f"  🤖 本文翻訳中: 全 {len(paragraphs)} 段落...")

    for i, para in enumerate(paragraphs):
        if not para.strip(): continue
        trans = translate_text_via_deepl(para)
        translated_paragraphs.append(trans)
        time.sleep(0.5) # レート制限対策
        if (i + 1) % 10 == 0:
            print(f"    ... {i + 1}/{len(paragraphs)} 完了")

    return "\n\n".join(translated_paragraphs)

# --- ▼ XML解析機能の強化 ▼ ---

def extract_data_from_xml(xml_content):
    """XMLからタイトル、本文、参考文献を抽出する"""
    try:
        root = ET.fromstring(xml_content)
        
        # 1. タイトル抽出
        title_node = root.find('.//tei:teiHeader//tei:titleStmt/tei:title', NAMESPACES)
        title = title_node.text.strip() if title_node is not None and title_node.text else "No Title Found"

        # 2. 本文抽出 (段落ごと)
        body_text_list = []
        paragraphs = root.findall('.//tei:text//tei:p', NAMESPACES)
        for p in paragraphs:
            # itertext()ですべてのタグ内のテキストを結合
            text = "".join(p.itertext()).strip()
            if text:
                body_text_list.append(text)
        full_body = "\n\n".join(body_text_list)

        # 3. 参考文献抽出
        references = []
        bib_structs = root.findall('.//tei:listBibl/tei:biblStruct', NAMESPACES)
        
        for i, bib in enumerate(bib_structs, 1):
            # 簡易的な抽出ロジック: タイトルと著者などを生のテキストとして結合
            # 本来は細かくタグをパースすべきだが、GROBIDの出力構造に合わせて簡易化
            ref_text_parts = []
            
            # タイトル (論文名 or 書籍名)
            ref_title = bib.find('.//tei:title', NAMESPACES)
            if ref_title is not None and ref_title.text:
                ref_text_parts.append(f"\"{ref_title.text}\"")
            
            # 発行年
            date = bib.find('.//tei:date', NAMESPACES)
            if date is not None and date.get('when'):
                ref_text_parts.append(f"({date.get('when')})")
            
            # 雑誌名など
            pub = bib.find('.//tei:publicationStmt/tei:publisher', NAMESPACES)
            if pub is not None and pub.text:
                ref_text_parts.append(pub.text)

            # もし構造化データがうまく取れなければ、noteタグなどを探す（簡易対応）
            full_ref_str = " ".join(ref_text_parts)
            if not full_ref_str:
                full_ref_str = "Extraction Failed"
            
            references.append(f"[{i}] {full_ref_str}")

        return {
            "title": title,
            "body": full_body,
            "references": references
        }

    except Exception as e:
        print(f"  ❌ XML解析エラー: {e}")
        return None

# --- ▼ Word生成機能 ▼ ---

def create_word_document(data, output_path):
    """翻訳結果をWordファイルとして保存"""
    doc = Document()

    # 1. タイトル (日本語 + 英語)
    doc.add_heading(data['jp_title'], 0) # 大きな見出し
    subtitle = doc.add_paragraph(data['en_title'])
    subtitle.italic = True # 原文タイトルは斜体で

    # 2. 本文 (日本語)
    doc.add_heading('本文 (Translated)', level=1)
    
    # 段落ごとにWordのパラグラフとして追加（読みやすさのため）
    paragraphs = data['jp_body'].split("\n\n")
    for p_text in paragraphs:
        p = doc.add_paragraph(p_text)
        p.paragraph_format.space_after = Pt(12) # 段落後の余白

    # 3. 参考文献 (原文のまま)
    if data['references']:
        doc.add_page_break() # 改ページ
        doc.add_heading('参考文献 (References)', level=1)
        for ref in data['references']:
            doc.add_paragraph(ref, style='List Number')

    doc.save(output_path)
    print(f"  💾 Word保存完了: {os.path.basename(output_path)}")

# --- ▼ メイン処理 ▼ ---

def process_single_pdf(pdf_path):
    base_filename = os.path.basename(pdf_path).replace(".pdf", "")
    xml_path = os.path.join(OUTPUT_XML_DIR, f"{base_filename}.xml")
    docx_path = os.path.join(OUTPUT_DOCX_DIR, f"{base_filename}_translated.docx")

    # 翻訳済み(docxが存在する)ならスキップ
    if os.path.exists(docx_path):
        print(f"\n⏭️  完全スキップ (完了済み): {base_filename}")
        return

    print(f"\n🔄 処理開始: {base_filename}")

    # 1. GROBID実行 & XML保存
    xml_content = ""
    if os.path.exists(xml_path):
        print("  📂 既存のXMLを使用します。")
        with open(xml_path, "r", encoding="utf-8") as f:
            xml_content = f.read()
    else:
        try:
            with open(pdf_path, 'rb') as f:
                files = {'input': f}
                resp = requests.post(GROBID_URL, files=files, data={'consolidateHeader': '1', 'consolidateCitations': '1'}, timeout=GROBID_TIMEOUT)
            
            if resp.status_code != 200:
                print(f"  ❌ GROBIDエラー: {resp.status_code}")
                return
            
            xml_content = resp.text
            with open(xml_path, "w", encoding="utf-8") as f:
                f.write(xml_content)
            print("  ✅ PDF解析完了 (GROBID)")
        except Exception as e:
            print(f"  ❌ GROBID接続エラー: {e}")
            return

    # 2. XMLからデータ抽出
    extracted_data = extract_data_from_xml(xml_content)
    if not extracted_data or not extracted_data['body']:
        print("  ⚠️ 本文抽出失敗")
        return

    # 3. 翻訳 (タイトルと本文)
    print(f"  🌍 タイトル翻訳中: {extracted_data['title'][:30]}...")
    jp_title = translate_text_via_deepl(extracted_data['title'])
    
    jp_body = translate_long_text(extracted_data['body'])

    # 4. Word生成用データ作成
    doc_data = {
        "en_title": extracted_data['title'],
        "jp_title": jp_title,
        "jp_body": jp_body,
        "references": extracted_data['references'] # 参考文献は翻訳しない
    }

    # 5. Word保存
    create_word_document(doc_data, docx_path)

def main():
    setup_directories()
    pdf_files = glob.glob(os.path.join(INPUT_DIR, "*.pdf"))
    
    if not pdf_files:
        print(f"'{INPUT_DIR}' にPDFがありません。")
        return

    print(f"--- {len(pdf_files)} 件のPDFをWord変換します ---")
    
    for pdf in pdf_files:
        process_single_pdf(pdf)

    print("\n--- 全ての処理が完了しました ---")

if __name__ == "__main__":
    main()